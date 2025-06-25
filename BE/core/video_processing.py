import os
import io
import json
import logging
import subprocess
import tempfile
from typing import Dict, Any, List, Optional

from BE.utils.helpers import ensure_dir_exists

try:
    from openai import OpenAI
except Exception:  # pragma: no cover - optional
    OpenAI = None

logger = logging.getLogger('ImageSearchApp')


# --- Metadata Extraction ---

def _run_ffprobe(video_path: str) -> Dict[str, Any]:
    """Run ffprobe and return parsed JSON."""
    cmd = [
        "ffprobe",
        "-v",
        "error",
        "-print_format",
        "json",
        "-show_format",
        "-show_streams",
        video_path,
    ]
    try:
        result = subprocess.run(cmd, capture_output=True, text=True, check=True)
        return json.loads(result.stdout)
    except Exception as e:  # pragma: no cover - external tool may not exist
        logger.error(f"ffprobe failed for {video_path}: {e}")
        return {}


def extract_comprehensive_metadata(video_path: str) -> Dict[str, Any]:
    """Extract basic and technical metadata from a video using ffprobe."""
    meta = {
        "success": False,
        "basic": {},
        "technical": {},
        "gps": {},
        "exif": {},
    }
    if not os.path.isfile(video_path):
        meta["error"] = "File not found"
        return meta

    data = _run_ffprobe(video_path)
    if not data:
        meta["error"] = "ffprobe_error"
        return meta

    try:
        format_info = data.get("format", {})
        streams = data.get("streams", [])
        video_stream = next((s for s in streams if s.get("codec_type") == "video"), {})
        audio_stream = next((s for s in streams if s.get("codec_type") == "audio"), {})

        meta["basic"] = {
            "filename": os.path.basename(video_path),
            "duration": float(format_info.get("duration", 0)),
            "size": int(format_info.get("size", 0)),
            "bit_rate": int(format_info.get("bit_rate", 0)),
            "format": format_info.get("format_long_name"),
            "video_codec": video_stream.get("codec_name"),
            "audio_codec": audio_stream.get("codec_name"),
            "width": video_stream.get("width"),
            "height": video_stream.get("height"),
        }

        meta["technical"] = {
            "frame_rate": eval(video_stream.get("r_frame_rate", "0")) if video_stream.get("r_frame_rate") else None,
            "color_space": video_stream.get("color_space"),
            "pixel_format": video_stream.get("pix_fmt"),
            "compression": video_stream.get("codec_long_name"),
        }

        tags = format_info.get("tags", {})
        gps_data = {}
        exif_data = {}
        for k, v in tags.items():
            if "location" in k.lower() or "gps" in k.lower():
                gps_data[k] = v
            else:
                exif_data[k] = v
        meta["gps"] = gps_data
        meta["exif"] = exif_data
        meta["success"] = True
    except Exception as e:  # pragma: no cover
        logger.error(f"Metadata parsing failed for {video_path}: {e}")
        meta["error"] = str(e)
    return meta


# --- Audio Transcription ---

def transcribe_with_timestamps(video_path: str) -> Dict[str, Any]:
    """Transcribe audio using Whisper with word-level timestamps."""
    if OpenAI is None:
        return {"success": False, "error": "OpenAI library not available"}

    if not os.path.isfile(video_path):
        return {"success": False, "error": "File not found"}

    try:
        client = OpenAI()
        with open(video_path, "rb") as f:
            transcript = client.audio.transcriptions.create(
                model="whisper-1",
                file=f,
                response_format="verbose_json",
                timestamp_granularities=["word"],
            )
        return {"success": True, "transcript": transcript}
    except Exception as e:  # pragma: no cover
        logger.error(f"Transcription failed for {video_path}: {e}")
        return {"success": False, "error": str(e)}


# --- Transcript Export ---

def _format_srt(words: List[Dict[str, Any]]) -> str:
    lines = []
    idx = 1
    for w in words:
        start = w.get("start", 0)
        end = w.get("end", 0)
        text = w.get("word", "").strip()
        if not text:
            continue
        start_s = _sec_to_srt_time(start)
        end_s = _sec_to_srt_time(end)
        lines.append(f"{idx}\n{start_s} --> {end_s}\n{text}\n")
        idx += 1
    return "\n".join(lines)


def _sec_to_srt_time(sec: float) -> str:
    h = int(sec // 3600)
    m = int((sec % 3600) // 60)
    s = int(sec % 60)
    ms = int((sec - int(sec)) * 1000)
    return f"{h:02}:{m:02}:{s:02},{ms:03}"


def export_transcript_multiple_formats(transcript_json: Dict[str, Any], output_dir: str, base_name: str) -> Dict[str, Any]:
    """Export transcript to multiple formats."""
    ensure_dir_exists(output_dir)
    result: Dict[str, Any] = {"success": False}
    words = []
    text_lines = []
    try:
        segments = transcript_json.get("segments", [])
        for seg in segments:
            for w in seg.get("words", []):
                words.append(w)
            text_lines.append(seg.get("text", "").strip())

        txt_path = os.path.join(output_dir, f"{base_name}.txt")
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write("\n".join(text_lines))

        json_path = os.path.join(output_dir, f"{base_name}.json")
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(transcript_json, f, ensure_ascii=False, indent=2)

        csv_path = os.path.join(output_dir, f"{base_name}.csv")
        with open(csv_path, "w", encoding="utf-8") as f:
            f.write("timestamp,speaker,text\n")
            for seg in segments:
                start = seg.get("start", 0)
                speaker = seg.get("speaker", "")
                text = seg.get("text", "").replace("\n", " ")
                f.write(f"{start},{speaker},{text}\n")

        srt_path = os.path.join(output_dir, f"{base_name}.srt")
        with open(srt_path, "w", encoding="utf-8") as f:
            f.write(_format_srt(words))

        docx_path = None
        try:
            from docx import Document  # type: ignore

            doc = Document()
            for seg in segments:
                start = _sec_to_srt_time(seg.get("start", 0))
                p = doc.add_paragraph()
                p.add_run(f"[{start}] ").bold = True
                p.add_run(seg.get("text", ""))
            docx_path = os.path.join(output_dir, f"{base_name}.docx")
            doc.save(docx_path)
        except Exception as e:  # pragma: no cover - docx optional
            logger.warning(f"DOCX export failed: {e}")

        result.update(
            {
                "success": True,
                "txt": txt_path,
                "json": json_path,
                "csv": csv_path,
                "srt": srt_path,
                "docx": docx_path,
            }
        )
    except Exception as e:  # pragma: no cover
        logger.error(f"Transcript export failed: {e}")
        result["error"] = str(e)
    return result


# --- Thumbnail Extraction ---

def extract_thumbnails(video_path: str, output_dir: str, interval: int = 5) -> List[str]:
    """Extract thumbnails at given interval (seconds) using ffmpeg."""
    ensure_dir_exists(output_dir)
    pattern = os.path.join(output_dir, "thumb_%04d.jpg")
    cmd = [
        "ffmpeg",
        "-y",
        "-i",
        video_path,
        "-vf",
        f"fps=1/{interval}",
        pattern,
    ]
    try:
        subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        files = sorted([os.path.join(output_dir, f) for f in os.listdir(output_dir) if f.startswith("thumb_")])
        return files
    except Exception as e:  # pragma: no cover
        logger.error(f"Thumbnail extraction failed for {video_path}: {e}")
        return []


# --- Report Generation ---

def generate_video_analysis_report(metadata: Dict[str, Any], transcript_json: Dict[str, Any], thumbnails: List[str], output_dir: str, base_name: str) -> Dict[str, Any]:
    """Generate a simple Word/Excel report with metadata and transcript."""
    ensure_dir_exists(output_dir)
    report_paths: Dict[str, Optional[str]] = {"success": False}
    try:
        docx_path = None
        xlsx_path = None
        try:
            from docx import Document  # type: ignore

            doc = Document()
            doc.add_heading("Video Analysis Report", 0)
            doc.add_heading("Metadata", level=1)
            for k, v in metadata.get("basic", {}).items():
                doc.add_paragraph(f"{k}: {v}")
            for k, v in metadata.get("technical", {}).items():
                doc.add_paragraph(f"{k}: {v}")
            doc.add_heading("Transcript", level=1)
            for seg in transcript_json.get("segments", []):
                ts = _sec_to_srt_time(seg.get("start", 0))
                doc.add_paragraph(f"[{ts}] {seg.get('text', '')}")
            docx_path = os.path.join(output_dir, f"{base_name}_report.docx")
            doc.save(docx_path)
        except Exception as e:  # pragma: no cover
            logger.warning(f"DOCX report failed: {e}")

        try:
            import xlsxwriter  # type: ignore

            xlsx_path = os.path.join(output_dir, f"{base_name}_report.xlsx")
            workbook = xlsxwriter.Workbook(xlsx_path)
            meta_sheet = workbook.add_worksheet("Metadata")
            row = 0
            for k, v in metadata.get("basic", {}).items():
                meta_sheet.write(row, 0, k)
                meta_sheet.write(row, 1, str(v))
                row += 1
            for k, v in metadata.get("technical", {}).items():
                meta_sheet.write(row, 0, k)
                meta_sheet.write(row, 1, str(v))
                row += 1
            trans_sheet = workbook.add_worksheet("Transcript")
            row = 0
            for seg in transcript_json.get("segments", []):
                trans_sheet.write(row, 0, seg.get("start"))
                trans_sheet.write(row, 1, seg.get("text"))
                row += 1
            workbook.close()
        except Exception as e:  # pragma: no cover
            logger.warning(f"XLSX report failed: {e}")

        report_paths.update({
            "success": True,
            "docx": docx_path,
            "xlsx": xlsx_path,
        })
    except Exception as e:  # pragma: no cover
        logger.error(f"Report generation failed: {e}")
        report_paths["error"] = str(e)
    return report_paths
